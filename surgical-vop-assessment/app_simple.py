"""
Vascular Anastomosis VOP Retrospective Assessment
Moderated v1: Evidence-based tri-state scoring (YES/NO/NULL) with coverage,
3-state proficiency, Economy Index, Item 6 split, and projection layer.
"""

import streamlit as st
import os
import re
import csv
import io
import json
import tempfile
import zipfile
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
from gemini_vision_client import (
    GeminiVisionClient,
    ITEM_LABELS,
    CHECKLIST_ITEMS,
    ECONOMY_ITEM,
    PROFICIENCY_ITEM,
    COMMENTS_ITEM,
    CSV_FIELDS,
    RED_LINE_LABELS,
    COACHING_TAGS_TAXONOMY,
    CORE_PROFICIENCY_ITEMS,
    build_result_json,
    extract_video_id,
    parse_candidate_ids_from_filename,
    get_canonical_output,
    project_to_docx,
)
from pdf_report import build_pdf_report, candidate_pdf_filename

RESULTS_DIR = Path(__file__).parent / "saved_results"
RESULTS_DIR.mkdir(parents=True, exist_ok=True)

VIDEOS_INBOX = Path(__file__).parent / "videos_inbox"
VIDEOS_INBOX.mkdir(parents=True, exist_ok=True)

# Default inbox folder. The user may point the app at any folder on disk in
# Inbox Folder mode; the chosen path is stored in session state and persisted
# in the queue file so Resume works across restarts.
DEFAULT_VIDEOS_INBOX = VIDEOS_INBOX


def get_inbox_root() -> Path:
    """Return the currently-selected inbox folder (defaults to videos_inbox/)."""
    sel = None
    try:
        sel = st.session_state.get("inbox_folder_path")
    except Exception:
        sel = None
    if sel:
        return Path(sel)
    return DEFAULT_VIDEOS_INBOX

QUEUE_FILE = RESULTS_DIR / "_queue.json"
VIDEO_EXTENSIONS = {".mp4", ".mov", ".avi", ".webm", ".mkv", ".m4v"}

st.set_page_config(
    page_title="Vascular Anastomosis VOP Assessment",
    page_icon="🩺",
    layout="wide",
)

MAX_FILE_SIZE_MB = 2000


def get_session_id() -> str:
    if "session_id" not in st.session_state:
        st.session_state.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    return st.session_state.session_id


def save_result_to_disk(result: Dict[str, Any], session_id: str):
    session_dir = RESULTS_DIR / session_id
    session_dir.mkdir(exist_ok=True)

    filename = result.get("filename", "unknown")
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in filename)
    result_file = session_dir / f"{safe_name}.json"

    with open(result_file, "w") as f:
        json.dump(result, f, indent=2, default=str)

    manifest_file = session_dir / "manifest.json"
    manifest = {
        "session_id": session_id,
        "created": datetime.now().isoformat(),
        "files": [],
    }
    if manifest_file.exists():
        with open(manifest_file) as f:
            manifest = json.load(f)

    if safe_name not in [Path(f).stem for f in manifest.get("files", [])]:
        manifest["files"].append(f"{safe_name}.json")
        manifest["updated"] = datetime.now().isoformat()

    with open(manifest_file, "w") as f:
        json.dump(manifest, f, indent=2)


def _safe_json_load(path) -> Optional[Dict[str, Any]]:
    """Load a JSON file defensively: never let a corrupt/oversized file crash the app."""
    try:
        with open(path) as f:
            return json.load(f)
    except RecursionError:
        # Deeply nested JSON can exceed the interpreter's recursion limit inside
        # Streamlit's script thread. Retry once with a temporarily raised limit.
        import sys
        old_limit = sys.getrecursionlimit()
        try:
            sys.setrecursionlimit(20000)
            with open(path) as f:
                return json.load(f)
        except Exception:
            return None
        finally:
            sys.setrecursionlimit(old_limit)
    except Exception:
        return None


def load_session_results(session_id: str) -> List[Dict[str, Any]]:
    session_dir = RESULTS_DIR / session_id
    if not session_dir.exists():
        return []

    results = []
    manifest_file = session_dir / "manifest.json"
    if manifest_file.exists():
        manifest = _safe_json_load(manifest_file) or {}
        skipped = []
        for filename in manifest.get("files", []):
            result_file = session_dir / filename
            if result_file.exists():
                data = _safe_json_load(result_file)
                if data is not None:
                    results.append(data)
                else:
                    skipped.append(filename)
        if skipped:
            st.warning(
                f"Skipped {len(skipped)} unreadable result file(s) in this session: "
                + ", ".join(skipped)
            )
    return results


def list_saved_sessions() -> List[Dict[str, Any]]:
    sessions = []
    if not RESULTS_DIR.exists():
        return sessions
    for session_dir in sorted(RESULTS_DIR.iterdir(), reverse=True):
        if session_dir.is_dir():
            manifest_file = session_dir / "manifest.json"
            if manifest_file.exists():
                manifest = _safe_json_load(manifest_file)
                if manifest is None:
                    continue
                sessions.append(
                    {
                        "session_id": session_dir.name,
                        "created": manifest.get("created", "Unknown"),
                        "file_count": len(manifest.get("files", [])),
                    }
                )
    return sessions


def build_batch_filename_stem(batch_results: List[Dict[str, Any]]) -> str:
    """Build a descriptive filename stem from batch results.

    Examples:
      Single video '2016 216-1116.m4v' → '2016_216-1116'
      Two videos '2020 2001.m4v', '2020 2005.m4v' → '2020_2001_2005'
      Mixed years → '2016_216-1116_2020_2001'
    """
    stems: List[str] = []
    for r in batch_results:
        fname = r.get("filename", "")
        if not fname:
            continue
        stem = Path(fname).stem.strip()
        normalized = re.sub(r'\s+', '_', stem)
        if normalized and normalized not in stems:
            stems.append(normalized)
    if not stems:
        return "batch"
    if len(stems) == 1:
        return stems[0]
    combined = "_".join(stems)
    if len(combined) > 120:
        combined = combined[:117] + "etc"
    return combined


def inbox_key(path: Path, root: Optional[Path] = None) -> str:
    """Stable queue/bookkeeping key for an inbox video: its path relative to
    the selected inbox root (POSIX form), so files in subfolders are unique
    and stable."""
    if root is None:
        root = get_inbox_root()
    try:
        return path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return path.name


def scan_inbox(root: Optional[Path] = None) -> List[Path]:
    """Recursively scan the selected inbox folder (including subfolders).

    Skips hidden directories and hidden files (names starting with '.').
    """
    if root is None:
        root = get_inbox_root()
    videos = []
    if root.exists() and root.is_dir():
        inbox_root = root.resolve()
        for f in sorted(root.rglob("*")):
            if not f.is_file():
                continue
            if f.suffix.lower() not in VIDEO_EXTENSIONS:
                continue
            # Skip hidden files or anything inside a hidden directory.
            try:
                rel_parts = f.resolve().relative_to(inbox_root).parts
            except ValueError:
                rel_parts = (f.name,)
            if any(part.startswith(".") for part in rel_parts):
                continue
            videos.append(f)
    return videos


def load_queue() -> Dict[str, Any]:
    if QUEUE_FILE.exists():
        try:
            with open(QUEUE_FILE) as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass
    return {"session_id": "", "videos": {}, "focused_mode": False}


def save_queue(queue: Dict[str, Any]):
    with open(QUEUE_FILE, "w") as f:
        json.dump(queue, f, indent=2, default=str)


def init_or_resume_queue(inbox_videos: List[Path], focused_mode: bool) -> Dict[str, Any]:
    queue = load_queue()

    existing_session = queue.get("session_id", "")
    completed_names = {
        name for name, info in queue.get("videos", {}).items()
        if info.get("status") == "completed"
    }

    inbox_root = get_inbox_root()

    if existing_session and completed_names:
        inbox_names = {inbox_key(v, inbox_root) for v in inbox_videos}
        new_names = inbox_names - set(queue.get("videos", {}).keys())
        if new_names:
            for video_path in inbox_videos:
                key = inbox_key(video_path, inbox_root)
                if key in new_names:
                    queue["videos"][key] = {
                        "path": str(video_path),
                        "status": "pending",
                        "added": datetime.now().isoformat(),
                    }
        # Keep the persisted folder in sync with the active selection.
        queue["video_folder"] = str(inbox_root)
        save_queue(queue)
        return queue

    session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    queue = {
        "session_id": session_id,
        "focused_mode": focused_mode,
        "created": datetime.now().isoformat(),
        "video_folder": str(inbox_root),
        "videos": {},
    }
    for video_path in inbox_videos:
        queue["videos"][inbox_key(video_path, inbox_root)] = {
            "path": str(video_path),
            "status": "pending",
            "added": datetime.now().isoformat(),
        }
    save_queue(queue)
    return queue


def process_video_from_path(
    video_path: str, filename: str, client: GeminiVisionClient, focused_mode: bool = False,
) -> Dict[str, Any]:
    try:
        file_size_mb = Path(video_path).stat().st_size / (1024 * 1024)
    except OSError as e:
        return {
            "filename": filename,
            "success": False,
            "error": f"Cannot access file: {e}",
            "timestamp": datetime.now().isoformat(),
        }

    if file_size_mb > MAX_FILE_SIZE_MB:
        return {
            "filename": filename,
            "success": False,
            "error": f"Video too large: {file_size_mb:.1f}MB (max {MAX_FILE_SIZE_MB}MB)",
            "timestamp": datetime.now().isoformat(),
        }

    try:
        if focused_mode:
            result = client.assess_video_items_1_3(video_path, filename=filename)
        else:
            result = client.assess_video(video_path, filename=filename)
        result["filename"] = filename
        result["timestamp"] = datetime.now().isoformat()
        return result
    except Exception as e:
        return {
            "filename": filename,
            "success": False,
            "error": str(e),
            "timestamp": datetime.now().isoformat(),
        }


def get_queue_summary(queue: Dict[str, Any]) -> Dict[str, int]:
    summary = {"pending": 0, "processing": 0, "completed": 0, "failed": 0, "total": 0}
    for info in queue.get("videos", {}).values():
        status = info.get("status", "pending")
        summary[status] = summary.get(status, 0) + 1
        summary["total"] += 1
    return summary


def load_completed_results_from_queue(queue: Dict[str, Any]) -> List[Dict[str, Any]]:
    results = []
    session_id = queue.get("session_id", "")
    if not session_id:
        return results
    session_dir = RESULTS_DIR / session_id
    if not session_dir.exists():
        return results
    for name, info in queue.get("videos", {}).items():
        if info.get("status") == "completed":
            result_file = info.get("result_file", "")
            if result_file:
                fpath = session_dir / result_file
                if fpath.exists():
                    data = _safe_json_load(fpath)
                    if data is not None:
                        results.append(data)
    return results


def clear_queue():
    if QUEUE_FILE.exists():
        os.unlink(str(QUEUE_FILE))


def init_session_state():
    if "batch_results" not in st.session_state:
        st.session_state.batch_results = []
    if "current_assessment" not in st.session_state:
        st.session_state.current_assessment = None
    if "processing" not in st.session_state:
        st.session_state.processing = False
    if "pending_files" not in st.session_state:
        st.session_state.pending_files = []
    if "current_index" not in st.session_state:
        st.session_state.current_index = 0
    if "total_files" not in st.session_state:
        st.session_state.total_files = 0
    if "surgeon_scores" not in st.session_state:
        st.session_state.surgeon_scores = {}
    if "surgeon_filename" not in st.session_state:
        st.session_state.surgeon_filename = ""
    if "focused_mode" not in st.session_state:
        st.session_state.focused_mode = False
    if "inbox_processing" not in st.session_state:
        st.session_state.inbox_processing = False
    if "inbox_queue" not in st.session_state:
        st.session_state.inbox_queue = None
    if "inbox_folder_path" not in st.session_state:
        # Restore a persisted custom folder from the queue file so Resume works
        # across restarts; otherwise default to videos_inbox/.
        persisted = load_queue().get("video_folder", "")
        st.session_state.inbox_folder_path = persisted or str(DEFAULT_VIDEOS_INBOX)
    if "inbox_auto_resumed" not in st.session_state:
        st.session_state.inbox_auto_resumed = False
        queue = load_queue()
        if queue.get("session_id"):
            for name, info in queue.get("videos", {}).items():
                if info.get("status") == "processing":
                    session_dir = RESULTS_DIR / queue["session_id"]
                    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
                    result_path = session_dir / f"{safe_name}.json"
                    if result_path.exists():
                        info["status"] = "completed"
                        info["result_file"] = f"{safe_name}.json"
                    else:
                        info["status"] = "pending"
            save_queue(queue)


def display_validation_error(result: Dict[str, Any], filename: str):
    st.error(f"Assessment validation failed for: {filename}")

    with st.expander("Why did this fail?", expanded=True):
        st.markdown("### Analysis")
        st.write(result.get("analysis", "Unknown issue"))

        st.markdown("### Advice")
        advice = result.get("advice", [])
        for item in advice:
            st.markdown(f"- {item}")

        if result.get("validation_details"):
            st.markdown("### Technical Details")
            details = result["validation_details"]
            st.write(f"- JSON parsed: {'Yes' if details.get('json_parsed') else 'No'}")
            st.write(f"- Candidates found: {details.get('candidate_count', 0)}")
            issues = details.get("issues", [])
            if issues:
                st.write("- Issues:")
                for issue in issues:
                    st.write(f"  - {issue}")

        if result.get("response"):
            st.markdown("### Raw Response (for debugging)")
            st.text_area("Gemini's response:", result["response"], height=200)


def get_item_by_id(items: List[Dict], item_id: int) -> Dict:
    for entry in items:
        if isinstance(entry, dict) and entry.get("item_id") == item_id:
            return entry
    return {}


def display_legacy_card(record: Dict[str, Any]):
    """Display a legacy (pre-2023) record without applying the new numbering."""
    case_id = record.get("case_id", "")
    legacy_rubric = record.get("rubric_version", "legacy")
    case_text = f" — Case: {case_id}" if case_id else ""
    st.warning(
        f"**Legacy rubric (pre-2023 refactor)**{case_text} — stored under the old "
        f"14-item numbering (`{legacy_rubric}`). This record is shown as-is and is "
        f"NOT reinterpreted under the current vop_2023_v1 numbering. The new stat "
        f"cards, proficiency/economy interpretation, and CSV export are disabled "
        f"for legacy records to avoid mislabeling."
    )
    items = record.get("items", [])
    if items:
        with st.expander("Raw legacy items (original numbering)", expanded=False):
            for it in items:
                if not isinstance(it, dict):
                    continue
                iid = it.get("item_id", "?")
                score = it.get("score", "--")
                ev = it.get("evidence", "")
                ev_str = ev if isinstance(ev, str) else json.dumps(ev, default=str)
                st.markdown(f"**Item {iid}:** {score}")
                if ev_str:
                    st.caption(ev_str[:300])


def build_action_plan(items: List[Dict[str, Any]]) -> List[tuple]:
    """Collect coaching tips into a prioritized action plan.

    Returns (item_id, label, coaching) tuples: core checklist NO items first,
    then remaining checklist NO items, then Economy (item 11) coaching when
    its score is below 5. Empty list when no tips exist.
    """
    plan = []
    for i in CHECKLIST_ITEMS:
        entry = get_item_by_id(items, i)
        coaching = entry.get("coaching", "")
        if entry.get("score") == "NO" and coaching:
            plan.append(
                (i, entry.get("label", ITEM_LABELS.get(i, f"Item {i}")), coaching)
            )
    plan.sort(key=lambda x: (0 if x[0] in CORE_PROFICIENCY_ITEMS else 1, x[0]))
    econ_entry = get_item_by_id(items, 11)
    econ_coaching = econ_entry.get("coaching", "")
    econ_score = econ_entry.get("score")
    if econ_coaching and isinstance(econ_score, int) and econ_score < 5:
        plan.append(
            (11, econ_entry.get("label", ITEM_LABELS.get(11, "Economy of Time and Motion")), econ_coaching)
        )
    return plan


def display_score_card(record: Dict[str, Any]):
    """Display the vop 2023 v1 score card with tri-state, coverage, projection toggle."""
    if record.get("is_legacy"):
        display_legacy_card(record)
        return
    items = record.get("items", [])
    proficiency = record.get("proficiency", "INSUFFICIENT_EVIDENCE")
    proficiency_rationale = record.get("proficiency_rationale", "")
    checklist_yes = record.get("checklist_yes_count", 0)
    null_count = record.get("null_count", 0)
    economy = record.get("economy_score")
    case_id = record.get("case_id", "")
    red_lines = record.get("red_lines_triggered", [])
    missing_core = record.get("missing_core_domains", [])
    coaching_tags = record.get("coaching_tags", [])
    coverage = record.get("coverage", {})

    if proficiency == "PROFICIENT":
        color = "green"
        border_color = "#4caf50"
        det_text = "PROFICIENT"
    elif proficiency == "NOT_PROFICIENT":
        color = "#f44336"
        border_color = "#f44336"
        det_text = "NOT PROFICIENT"
    else:
        color = "#ff9800"
        border_color = "#ff9800"
        det_text = "INSUFFICIENT EVIDENCE"

    cov_pct = coverage.get("observed_percent", 0)
    obs_count = coverage.get("observed_count", 0)
    null_text = f" ({null_count} not observed)" if null_count > 0 else ""
    case_text = f"<span style='color:#90a4ae;font-size:0.8em;'>Case: {case_id}</span><br/>" if case_id else ""
    econ_display = economy if isinstance(economy, int) else "--"

    st.markdown(
        f"""
        <div style="background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
                    border-radius: 12px; padding: 20px; margin: 10px 0;
                    border-left: 4px solid {border_color};">
            {case_text}
            <h3 style="margin:0; color: #e0e0e0;">Checklist: {checklist_yes} / 10 YES{null_text} &nbsp;&nbsp;|&nbsp;&nbsp; Economy: {econ_display} / 5 &nbsp;&nbsp;|&nbsp;&nbsp; Coverage: {cov_pct:.0f}%</h3>
            <p style="margin:5px 0 0 0; color: {color}; font-weight: bold; font-size: 1.2em;">
                {det_text}
            </p>
            <p style="margin:3px 0 0 0; color: #b0b0b0; font-size: 0.85em;">
                {proficiency_rationale}
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    action_plan = build_action_plan(items)
    if action_plan:
        st.markdown("#### Action Plan")
        plan_rows = []
        for iid, label, coaching in action_plan:
            core_badge = (
                " <span style='color:#ef5350;font-size:0.75em;'>[core]</span>"
                if iid in CORE_PROFICIENCY_ITEMS
                else ""
            )
            plan_rows.append(
                f"<div style='margin:4px 0;color:#e0e0e0;'>"
                f"<b>Item {iid} — {label}</b>{core_badge}<br/>"
                f"<span style='color:#4fc3f7;font-size:0.9em;'>&#8594; {coaching}</span>"
                f"</div>"
            )
        st.markdown(
            "<div style='background:#16213e;border-radius:8px;padding:12px 16px;"
            "border-left:4px solid #4fc3f7;margin:6px 0 10px 0;'>"
            + "".join(plan_rows)
            + "</div>",
            unsafe_allow_html=True,
        )

    core_obs = coverage.get("core_observed", {})
    if core_obs:
        core_badges = []
        for iid in ["2", "7", "9", "10"]:
            observed = core_obs.get(iid, False)
            badge_color = "#2e7d32" if observed else "#616161"
            badge_label = f"Item {iid}: {'observed' if observed else 'not observed'}"
            core_badges.append(
                f"<span style='background:{badge_color};color:white;font-size:0.7em;padding:2px 6px;border-radius:3px;margin-right:4px;'>{badge_label}</span>"
            )
        st.markdown("**Core Domains:** " + " ".join(core_badges), unsafe_allow_html=True)

    if red_lines:
        st.markdown("#### Red-Line Failures")
        for rl in red_lines:
            st.markdown(
                f"<span style='color:#f44336;font-weight:bold;'>&#9888; {rl}</span>",
                unsafe_allow_html=True,
            )

    if missing_core:
        st.markdown("#### Missing Core Domains")
        for mc in missing_core:
            st.markdown(
                f"<span style='color:#ff9800;'>&#9888; {mc}</span>",
                unsafe_allow_html=True,
            )

    if coaching_tags:
        st.markdown("#### Coaching Tags")
        tag_html = " ".join(
            f"<span style='background:#37474f;color:#b0bec5;font-size:0.8em;padding:2px 8px;border-radius:4px;margin-right:4px;'>{tag}</span>"
            for tag in coaching_tags
        )
        st.markdown(tag_html, unsafe_allow_html=True)

    st.markdown("#### Procedural Checklist (Items 1-10)")
    for i in CHECKLIST_ITEMS:
        entry = get_item_by_id(items, i)
        score = entry.get("score", "--")
        observability = entry.get("observability", "OBSERVED")
        label = entry.get("label", ITEM_LABELS.get(i, f"Item {i}"))
        evidence = entry.get("evidence", "")

        if score == "YES":
            icon = "&#10004;"
            score_color = "#4caf50"
            score_display = "YES"
            obs_badge = "<span style='background:#2e7d32;color:white;font-size:0.7em;padding:1px 5px;border-radius:3px;margin-left:5px;'>observed</span>"
        elif score == "NO":
            icon = "&#10008;"
            score_color = "#f44336"
            score_display = "NO"
            obs_badge = "<span style='background:#c62828;color:white;font-size:0.7em;padding:1px 5px;border-radius:3px;margin-left:5px;'>observed</span>"
        elif score == "NULL":
            icon = "&#8212;"
            score_color = "#78909c"
            score_display = "NULL"
            obs_badge = "<span style='background:#616161;color:white;font-size:0.7em;padding:1px 5px;border-radius:3px;margin-left:5px;'>not observed</span>"
        else:
            icon = "—"
            score_color = "#9e9e9e"
            score_display = str(score)
            obs_badge = ""

        non_gating_tag = ""
        if i == 6:
            non_gating_tag = " <span style='color:#90a4ae;font-size:0.7em;'>[non-gating]</span>"
        core_tag = ""
        if i in CORE_PROFICIENCY_ITEMS:
            core_tag = " <span style='color:#ef5350;font-size:0.7em;'>[core]</span>"

        st.markdown(
            f"**{i}. {label}:** <span style='color:{score_color};font-weight:bold;'>{icon} {score_display}</span>{obs_badge}{non_gating_tag}{core_tag}",
            unsafe_allow_html=True,
        )

        if i == 6:
            subitems = entry.get("subitems", {})
            sub_6a = subitems.get("6a_right_angle_method", {})
            sub_6b = subitems.get("6b_safe_transfer_outcome", {})
            if sub_6a or sub_6b:
                a_score = sub_6a.get("score", "--")
                b_score = sub_6b.get("score", "--")
                a_color = {"YES": "#4caf50", "NO": "#f44336", "NULL": "#78909c"}.get(a_score, "#9e9e9e")
                b_color = {"YES": "#4caf50", "NO": "#f44336", "NULL": "#78909c"}.get(b_score, "#9e9e9e")
                st.markdown(
                    f"&nbsp;&nbsp;&nbsp;&nbsp;6a Method: <span style='color:{a_color};font-weight:bold;'>{a_score}</span> &nbsp;|&nbsp; "
                    f"6b Outcome: <span style='color:{b_color};font-weight:bold;'>{b_score}</span>",
                    unsafe_allow_html=True,
                )
                if sub_6a.get("evidence"):
                    st.caption(f"  6a: {sub_6a['evidence']}")
                if sub_6b.get("evidence"):
                    st.caption(f"  6b: {sub_6b['evidence']}")

        if evidence and i != 6:
            if isinstance(evidence, str):
                st.caption(f"Evidence: {evidence}")
            elif isinstance(evidence, dict):
                st.caption(f"Evidence: {json.dumps(evidence, indent=2)}")

        coaching = entry.get("coaching", "")
        if score == "NO" and coaching:
            st.markdown(
                f"<div style='color:#4fc3f7;font-size:0.9em;margin:2px 0 8px 0;'>&#8594; <b>Coaching:</b> {coaching}</div>",
                unsafe_allow_html=True,
            )
        elif score == "NULL":
            st.markdown(
                "<div style='color:#9e9e9e;font-size:0.85em;font-style:italic;margin:2px 0 8px 0;'>Not observed on video — no coaching given</div>",
                unsafe_allow_html=True,
            )

    st.markdown("---")

    st.markdown("#### Economy of Time and Motion (Item 11)")
    econ_entry = get_item_by_id(items, 11)
    econ_score = econ_entry.get("score", "--")
    econ_evidence = econ_entry.get("evidence", "")
    econ_obs = econ_entry.get("observability", "OBSERVED")

    if econ_score == "NULL" or econ_obs == "NOT_OBSERVED":
        st.markdown(
            f"**Score:** <span style='color:#78909c;font-weight:bold;font-size:1.3em;'>NULL</span> "
            "<span style='background:#616161;color:white;font-size:0.7em;padding:1px 5px;border-radius:3px;margin-left:5px;'>not observed</span> "
            "<span style='color:#b0b0b0;font-size:0.85em;'>(does not gate proficiency)</span>",
            unsafe_allow_html=True,
        )
    else:
        econ_colors = {1: "#f44336", 2: "#ff9800", 3: "#ffc107", 4: "#8bc34a", 5: "#4caf50"}
        econ_color = econ_colors.get(econ_score if isinstance(econ_score, int) else 0, "#9e9e9e")
        st.markdown(
            f"**Score:** <span style='color:{econ_color};font-weight:bold;font-size:1.3em;'>{econ_score} / 5</span> <span style='color:#b0b0b0;font-size:0.85em;'>(out of 5; does not gate proficiency)</span>",
            unsafe_allow_html=True,
        )

    if isinstance(econ_evidence, dict):
        flow_org = econ_evidence.get("flow_organization", "")
        wasted_events = econ_evidence.get("wasted_motion_events", [])
        ei = econ_evidence.get("economy_index")

        if flow_org and flow_org != "unknown":
            flow_colors = {"organized": "#4caf50", "mixed": "#ffc107", "disorganized": "#f44336"}
            fc = flow_colors.get(flow_org, "#9e9e9e")
            st.markdown(f"**Flow organization:** <span style='color:{fc};font-weight:bold;'>{flow_org}</span>", unsafe_allow_html=True)

        if ei is not None:
            st.markdown(f"**Economy Index (EI):** {ei:.1f}")

        if wasted_events and isinstance(wasted_events, list) and len(wasted_events) > 0:
            st.markdown(f"**Wasted-motion events ({len(wasted_events)}):**")
            for idx, evt in enumerate(wasted_events, 1):
                if isinstance(evt, dict):
                    evt_type = evt.get("type", "other")
                    note = evt.get("note", "")
                    count = evt.get("count_estimate", 1)
                    weight = {"instrument_search": 2, "pause_reset": 2, "failed_pass_sequence": 2, "excess_regrasp_cluster": 1, "other": 1}.get(evt_type, 1)
                    st.caption(f"{idx}. {evt_type} (x{int(count)}, wt={weight}): {note}" if note else f"{idx}. {evt_type} (x{int(count)}, wt={weight})")
                else:
                    st.caption(f"{idx}. {evt}")
    elif isinstance(econ_evidence, str) and econ_evidence:
        st.caption(f"Evidence: {econ_evidence}")

    econ_coaching = econ_entry.get("coaching", "")
    if econ_coaching:
        if econ_score == 5:
            st.markdown(
                f"<div style='color:#66bb6a;font-size:0.9em;margin:4px 0 8px 0;'>&#9733; <b>{econ_coaching}</b></div>",
                unsafe_allow_html=True,
            )
        elif isinstance(econ_score, int):
            st.markdown(
                f"<div style='color:#4fc3f7;font-size:0.9em;margin:4px 0 8px 0;'>&#8594; <b>Coaching:</b> {econ_coaching}</div>",
                unsafe_allow_html=True,
            )

    diagnostics = record.get("diagnostics", {})
    econ_warnings = diagnostics.get("economy_tally_warnings", [])
    if econ_warnings:
        for w in econ_warnings:
            st.warning(w)

    st.markdown("---")

    st.markdown("#### Final Proficiency (Item 12)")
    st.markdown(
        f"**Proficiency:** <span style='color:{color};font-weight:bold;'>{det_text}</span>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<span style='color:#90a4ae;font-size:0.85em;'>Core domains: Items 2, 7, 9, 10. "
        "NOT_PROFICIENT if any core NO. INSUFFICIENT_EVIDENCE if any core NULL. "
        "Item 6 and Item 11 are non-gating.</span>",
        unsafe_allow_html=True,
    )
    if proficiency_rationale:
        st.caption(f"Rationale: {proficiency_rationale}")

    comments_entry = get_item_by_id(items, 13)
    comments = comments_entry.get("score", "")
    item_13_evidence = comments_entry.get("evidence", {})
    if isinstance(item_13_evidence, dict):
        item_13_tags = item_13_evidence.get("coaching_tags", [])
    else:
        item_13_tags = comments_entry.get("coaching_tags", [])
    if comments:
        st.markdown("---")
        st.markdown("#### Summative Comments (Item 13)")
        st.write(comments)
        if item_13_tags:
            tag_html = " ".join(
                f"<span style='background:#263238;color:#80cbc4;font-size:0.75em;padding:2px 6px;border-radius:3px;margin-right:3px;'>{t}</span>"
                for t in item_13_tags
            )
            st.markdown(tag_html, unsafe_allow_html=True)

    projected = record.get("projected_docx", {})
    if projected:
        with st.expander("DOCX-Compatible Projection"):
            proj_mode = projected.get("mode", "construct_mode")
            st.markdown(f"**Projection Mode:** `{proj_mode}`")

            proj_items = projected.get("items_1_10", projected.get("items_1_11", {}))
            proj_econ = projected.get("economy_1_5", projected.get("economy_1_4", "--"))
            proj_prof = projected.get("proficiency_yes_no", "--")
            proj_notes = projected.get("projection_notes", [])

            proj_yes = sum(1 for v in proj_items.values() if v == "YES")
            st.markdown(f"**Projected Checklist:** {proj_yes}/10 YES | Economy: {proj_econ}/5 | Proficiency: {proj_prof}")

            for k, v in sorted(proj_items.items(), key=lambda x: int(x[0])):
                pc = "#4caf50" if v == "YES" else "#f44336"
                st.markdown(f"Item {k}: <span style='color:{pc};font-weight:bold;'>{v}</span>", unsafe_allow_html=True)

            if proj_notes:
                st.markdown("**Projection Notes:**")
                for n in proj_notes:
                    st.caption(f"- {n}")

    for_data = record.get("for_data")
    if for_data:
        with st.expander("Factual Observation Record (FOR)"):
            st.json(for_data)

    diagnostics = record.get("diagnostics", {})
    scoring_warnings = diagnostics.get("scoring_warnings", [])
    scoring_path = record.get("scoring_path", "")
    if scoring_path:
        st.caption(f"Scoring path: {scoring_path}")
    if scoring_warnings:
        with st.expander("Scoring Diagnostics"):
            for w in scoring_warnings:
                st.warning(w)


def process_single_video(
    uploaded_file, client: GeminiVisionClient, focused_mode: bool = False,
) -> Dict[str, Any]:
    file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)

    if file_size_mb > MAX_FILE_SIZE_MB:
        return {
            "filename": uploaded_file.name,
            "success": False,
            "error": f"Video too large: {file_size_mb:.1f}MB (max {MAX_FILE_SIZE_MB}MB)",
            "timestamp": datetime.now().isoformat(),
        }

    with tempfile.NamedTemporaryFile(
        delete=False, suffix=Path(uploaded_file.name).suffix
    ) as tmp:
        tmp.write(uploaded_file.getvalue())
        temp_path = tmp.name

    try:
        if focused_mode:
            result = client.assess_video_items_1_3(temp_path, filename=uploaded_file.name)
        else:
            result = client.assess_video(temp_path, filename=uploaded_file.name)
        result["filename"] = uploaded_file.name
        result["timestamp"] = datetime.now().isoformat()
        return result
    finally:
        if os.path.exists(temp_path):
            os.unlink(temp_path)


def display_focused_result(result: Dict[str, Any]):
    """Display Items 1 & 3 focused assessment result."""
    filename = result.get("filename", "Unknown")

    if not result.get("success"):
        error_msg = result.get('error', 'Unknown error')
        st.error(f"Error processing {filename}: {error_msg}")
        return

    scored_outputs = result.get("scored_outputs", [])
    if not scored_outputs:
        st.warning(f"No scored outputs for {filename}")
        return

    for scored in scored_outputs:
        items = scored.get("items", [])
        case_id = scored.get("case_id", filename)

        st.subheader(f"Items 1 & 3 — {case_id}")

        for item in items:
            iid = item.get("item_id")
            score = item.get("score", "NULL")
            obs = item.get("observability", "NOT_OBSERVED")
            evidence = item.get("evidence", "")
            label = ITEM_LABELS.get(iid, f"Item {iid}")

            if score == "YES":
                icon = "✅"
                color = "#4caf50"
            elif score == "NO":
                icon = "❌"
                color = "#f44336"
            else:
                icon = "⬜"
                color = "#9e9e9e"

            st.markdown(
                f"<div style='background:#222;border-radius:8px;padding:12px;margin:8px 0;'>"
                f"<span style='font-size:1.2em;'>{icon}</span>&nbsp;&nbsp;"
                f"<strong>Item {iid}</strong>: {label}&nbsp;&nbsp;"
                f"<span style='color:{color};font-weight:bold;font-size:1.1em;'>{score}</span>"
                f"&nbsp;&nbsp;<span style='color:#888;'>({obs})</span>"
                f"</div>",
                unsafe_allow_html=True,
            )
            if evidence:
                st.caption(f"Evidence: {evidence}")

        st.markdown("---")

    warnings = result.get("scoring_warnings", [])
    if warnings:
        with st.expander("Scoring Warnings"):
            for w in warnings:
                st.warning(w)


def generate_focused_csv(batch_results: List[Dict[str, Any]]) -> str:
    """Generate CSV with only Items 1 & 3 scores.

    Column names match the full CSV (case_id, item_1, item_1_observability,
    item_3, item_3_observability) so a statistician can merge/vlookup on
    case_id and directly replace columns in the original results.
    Extra _evidence columns are appended for reference.
    """
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "case_id", "video_id",
        "item_1", "item_1_observability", "item_1_evidence",
        "item_3", "item_3_observability", "item_3_evidence",
    ])

    for result in batch_results:
        if not result.get("success"):
            continue
        fname = result.get("filename", "")
        scored_outputs = result.get("scored_outputs", [])
        for scored in scored_outputs:
            case_id = scored.get("case_id", fname)
            video_id = Path(fname).stem if fname else ""
            items = scored.get("items", [])
            item1 = next((i for i in items if isinstance(i, dict) and i.get("item_id") == 1), {})
            item3 = next((i for i in items if isinstance(i, dict) and i.get("item_id") == 3), {})
            writer.writerow([
                case_id, video_id,
                item1.get("score", ""), item1.get("observability", ""), item1.get("evidence", ""),
                item3.get("score", ""), item3.get("observability", ""), item3.get("evidence", ""),
            ])

    return output.getvalue()


def display_assessment_result(result: Dict[str, Any]):
    filename = result.get("filename", "Unknown")

    if result.get("success"):
        records = build_result_json(
            filename, result["response"], result.get("timestamp", ""),
            scored_outputs=result.get("scored_outputs"),
            scoring_warnings=result.get("scoring_warnings"),
        )

        num_candidates = len(records)
        if num_candidates > 1:
            st.success(f"Assessment complete: {filename} — {num_candidates} candidates detected")
        else:
            st.success(f"Assessment complete: {filename}")

        for rec_idx, record in enumerate(records):
            candidate_id = record.get("candidate_id", "")
            segment_start = record.get("segment_start", "")

            if num_candidates > 1:
                header = f"Candidate {candidate_id}"
                if segment_start:
                    header += f" (starts at {segment_start})"
                st.markdown(f"### {header}")

            display_score_card(record)

            st.markdown("---")

            col_j1, col_j2, col_j3, col_j4 = st.columns(4)
            with col_j1:
                canonical = get_canonical_output(record)
                canonical_str = json.dumps(canonical, indent=2)
                json_filename = Path(filename).stem + f"_candidate_{candidate_id}.json" if num_candidates > 1 else Path(filename).stem + ".json"
                st.download_button(
                    f"Download Canonical JSON",
                    data=canonical_str,
                    file_name=json_filename,
                    mime="application/json",
                    key=f"canonical_{filename}_{candidate_id}_{result.get('timestamp', '')}",
                )
            with col_j2:
                full_str = json.dumps(record, indent=2, default=str)
                st.download_button(
                    f"Download Full Record",
                    data=full_str,
                    file_name=f"vop_full_{Path(filename).stem}_{candidate_id}.json",
                    mime="application/json",
                    key=f"full_{filename}_{candidate_id}_{result.get('timestamp', '')}",
                )
            with col_j3:
                if rec_idx == 0:
                    st.download_button(
                        f"Download Raw FOR",
                        data=result["response"],
                        file_name=f"vop_for_{Path(filename).stem}.txt",
                        mime="text/plain",
                        key=f"for_{filename}_{result.get('timestamp', '')}",
                    )
            with col_j4:
                if not record.get("is_legacy"):
                    try:
                        pdf_bytes = build_pdf_report(record)
                        st.download_button(
                            "Download PDF report",
                            data=pdf_bytes,
                            file_name=candidate_pdf_filename(record),
                            mime="application/pdf",
                            key=f"pdf_{filename}_{candidate_id}_{result.get('timestamp', '')}",
                        )
                    except Exception as e:
                        st.caption(f"PDF unavailable: {e}")

            if num_candidates > 1 and rec_idx < num_candidates - 1:
                st.markdown("---")
                st.markdown("---")

    elif result.get("validation_failed"):
        display_validation_error(result, filename)
    else:
        error_msg = result.get('error', 'Unknown error')
        error_lower = error_msg.lower()
        if any(k in error_lower for k in ["429", "resource_exhausted", "rate limit", "quota", "exhausted"]):
            st.error(
                f"Rate limit exceeded for {filename}. The AI service was too busy even after "
                f"multiple retries. Please wait a few minutes and try again, or process fewer videos at once."
            )
            st.caption(f"Technical detail: {error_msg}")
        else:
            st.error(
                f"Error processing {filename}: {error_msg}"
            )


def generate_scores_csv(batch_results: List[Dict[str, Any]]) -> str:
    output = io.StringIO()

    successful_results = [r for r in batch_results if r.get("success")]
    if not successful_results:
        return ""

    headers = ["case_id", "video_id", "candidate_id", "segment_start"]
    for i in range(1, 11):
        headers.append(f"item_{i}")
        headers.append(f"item_{i}_observability")
        if i == 6:
            headers.append("item_6a_method")
            headers.append("item_6b_outcome")
        headers.append(f"item_{i}_coaching")
    headers += [
        "item_11_economy", "item_11_flow_organization",
        "item_11_economy_index", "item_11_wasted_events_count", "item_11_coaching",
        "item_12_proficient", "proficiency_rationale",
        "observed_count", "observed_percent", "core_observed",
        "red_lines", "missing_core_domains", "coaching_tags", "comments",
    ]
    writer = csv.writer(output)
    writer.writerow(headers)

    rows_written = 0
    for result in successful_results:
        filename = result.get("filename", "Unknown")
        response = result.get("response", "")

        records = build_result_json(
            filename, response, result.get("timestamp", ""),
            scored_outputs=result.get("scored_outputs"),
            scoring_warnings=result.get("scoring_warnings"),
        )
        for record in records:
            if record.get("is_legacy"):
                # Legacy (pre-2023) record: do not export under new numbering.
                continue
            rows_written += 1
            items = record.get("items", [])
            row = [
                record.get("case_id", ""),
                record.get("video_id", ""),
                record.get("candidate_id", ""),
                record.get("segment_start", ""),
            ]

            for i in range(1, 11):
                entry = get_item_by_id(items, i)
                row.append(entry.get("score", ""))
                row.append(entry.get("observability", ""))
                if i == 6:
                    subitems = entry.get("subitems", {})
                    sub_6a = subitems.get("6a_right_angle_method", {})
                    sub_6b = subitems.get("6b_safe_transfer_outcome", {})
                    row.append(sub_6a.get("score", ""))
                    row.append(sub_6b.get("score", ""))
                row.append(entry.get("coaching", ""))

            econ_entry = get_item_by_id(items, 11)
            row.append(str(econ_entry.get("score", "")) if econ_entry else "")

            econ_evidence = econ_entry.get("evidence", {}) if econ_entry else {}
            if isinstance(econ_evidence, dict):
                row.append(econ_evidence.get("flow_organization", ""))
                row.append(str(econ_evidence.get("economy_index", "")) if econ_evidence.get("economy_index") is not None else "")
                events = econ_evidence.get("wasted_motion_events", [])
                row.append(str(len(events)) if isinstance(events, list) else "")
            else:
                row.extend(["", "", ""])
            row.append(econ_entry.get("coaching", "") if econ_entry else "")

            row.append(record.get("proficiency", ""))
            row.append(record.get("proficiency_rationale", ""))

            coverage = record.get("coverage", {})
            row.append(str(coverage.get("observed_count", "")))
            row.append(f"{coverage.get('observed_percent', 0):.0f}")
            core_obs = coverage.get("core_observed", {})
            core_str = "; ".join(f"{k}={'Y' if v else 'N'}" for k, v in sorted(core_obs.items())) if core_obs else ""
            row.append(core_str)

            rl = record.get("red_lines_triggered", [])
            row.append("; ".join(rl) if rl else "None")

            mc = record.get("missing_core_domains", [])
            row.append("; ".join(mc) if mc else "")

            ct = record.get("coaching_tags", [])
            if not ct:
                item_13 = get_item_by_id(record.get("items", []), 13)
                if item_13:
                    evidence_13 = item_13.get("evidence", {})
                    if isinstance(evidence_13, dict):
                        ct = evidence_13.get("coaching_tags", [])
            row.append("; ".join(ct) if ct else "")

            comments_entry = get_item_by_id(items, 13)
            c = comments_entry.get("score", "") if comments_entry else ""
            row.append(c)

            writer.writerow(row)

    if rows_written == 0:
        return ""
    return output.getvalue()


def generate_batch_json(batch_results: List[Dict[str, Any]]) -> str:
    successful_results = [r for r in batch_results if r.get("success")]
    if not successful_results:
        return ""

    all_records = []
    for result in successful_results:
        records = build_result_json(
            result.get("filename", "Unknown"),
            result.get("response", ""),
            result.get("timestamp", ""),
            scored_outputs=result.get("scored_outputs"),
            scoring_warnings=result.get("scoring_warnings"),
        )
        all_records.extend(records)

    canonical = [get_canonical_output(r) for r in all_records]
    return json.dumps(canonical, indent=2)


def generate_all_pdfs_zip(batch_results: List[Dict[str, Any]]) -> bytes:
    """Build a ZIP archive containing one PDF report per (non-legacy) candidate.

    Returns empty bytes if there is nothing to export.
    """
    successful_results = [r for r in batch_results if r.get("success")]
    if not successful_results:
        return b""

    buf = io.BytesIO()
    used_names: Dict[str, int] = {}
    wrote_any = False
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for result in successful_results:
            records = build_result_json(
                result.get("filename", "Unknown"),
                result.get("response", ""),
                result.get("timestamp", ""),
                scored_outputs=result.get("scored_outputs"),
                scoring_warnings=result.get("scoring_warnings"),
            )
            for record in records:
                if record.get("is_legacy"):
                    continue
                try:
                    pdf_bytes = build_pdf_report(record)
                except Exception:
                    continue
                name = candidate_pdf_filename(record)
                # De-duplicate archive entry names.
                if name in used_names:
                    used_names[name] += 1
                    stem = name[:-4] if name.endswith(".pdf") else name
                    name = f"{stem}_{used_names[name]}.pdf"
                else:
                    used_names[name] = 0
                zf.writestr(name, pdf_bytes)
                wrote_any = True

    return buf.getvalue() if wrote_any else b""


def parse_docx_surgeon_scores(file_bytes: bytes, filename: str = "") -> Dict[str, Dict[str, str]]:
    """Parse surgeon scores from a VOP DOCX assessment form.

    Supports three layouts:
      1. Column-per-item table: header row with item numbers (1-10), data rows per candidate
      2. Row-per-item table: each table row has one item number + YES/NO score
      3. Paragraph-based: sequential paragraphs with bold item numbers, descriptions, and YES/NO

    Returns dict keyed by case_id/candidate_id with values like:
      {"item_1": "YES", ..., "item_10": "NO", "item_11_economy": "3", "item_12_proficient": "YES"}
    """
    import re
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    all_tables: List[List[List[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            all_tables.append(rows)

    full_text = "\n".join(p.text for p in doc.paragraphs)

    candidate_ids: List[str] = []
    id_label_pat = re.compile(
        r"(?:candidate|case)[\s_\-:/#]*(\d{4}[\s_\-]?\d{3,4})", re.IGNORECASE
    )
    matches = id_label_pat.findall(full_text)
    if matches:
        candidate_ids.extend(m.replace(" ", "").replace("-", "_") for m in matches)
    else:
        id_short_pat = re.compile(
            r"(?:candidate|case)[\s_\-:/#]*(\d{3,6})", re.IGNORECASE
        )
        candidate_ids.extend(id_short_pat.findall(full_text))

    if not candidate_ids and filename:
        fn_base = filename.rsplit(".", 1)[0] if "." in filename else filename
        fn_id_m = re.match(r"(\d{4})[\s_\-]+(\d{3,4})", fn_base)
        if fn_id_m:
            candidate_ids.append(f"{fn_id_m.group(1)}_{fn_id_m.group(2)}")
        else:
            fn_id_m2 = re.match(r"(\d{4}[\s_\-]+\d{3,4})", fn_base)
            if fn_id_m2:
                candidate_ids.append(re.sub(r"[\s\-]+", "_", fn_id_m2.group(1)))

    for tbl in all_tables:
        for row_cells in tbl:
            combined = " ".join(row_cells).lower()
            if "candidate" in combined or "case" in combined:
                for cell in row_cells:
                    for m in id_label_pat.finditer(cell):
                        cid = m.group(1).replace(" ", "").replace("-", "_")
                        if cid not in candidate_ids:
                            candidate_ids.append(cid)
                    if not id_label_pat.search(cell):
                        for m2 in re.finditer(r"(?:candidate|case)[\s_\-:/#]*(\d{3,6})", cell, re.IGNORECASE):
                            cid = m2.group(1)
                            if cid not in candidate_ids:
                                candidate_ids.append(cid)

    yes_no_re = re.compile(r"^(yes|no|y|n)$", re.IGNORECASE)
    score_1_5_re = re.compile(r"^([1-5])$")
    prof_re = re.compile(r"\b(not\s*proficient|proficient|yes|no)\b", re.IGNORECASE)

    def _normalize_yn(val: str) -> str:
        v = val.strip().upper()
        return "YES" if v in ("YES", "Y") else "NO"

    def _normalize_prof(val: str) -> str:
        v = val.strip().upper()
        if "NOT" in v:
            return "NO"
        if v in ("YES", "PROFICIENT"):
            return "YES"
        return "NO"

    def _has_yes_no(val: str) -> bool:
        return bool(yes_no_re.match(val.strip()))

    surgeon_rows: Dict[str, Dict[str, str]] = {}

    for tbl in all_tables:
        if len(tbl) < 2:
            continue

        header = tbl[0]
        header_lower = [h.lower().strip() for h in header]

        item_col_map: Dict[int, int] = {}
        econ_col = -1
        prof_col = -1
        case_col = -1

        item_hdr_re = re.compile(r"(?:item\s*)?#?\s*(\d{1,2})\s*$", re.IGNORECASE)

        for ci, h in enumerate(header_lower):
            if any(kw in h for kw in ["case", "candidate"]) and "id" in h or h in ("case", "candidate", "id", "case_id", "candidate_id"):
                case_col = ci
            elif "economy" in h or "item 11" in h or h == "11":
                econ_col = ci
            elif "proficien" in h or "item 12" in h or h == "12":
                prof_col = ci
            else:
                m = item_hdr_re.search(h)
                if m:
                    num = int(m.group(1))
                    if 1 <= num <= 10:
                        item_col_map[num] = ci

        if len(item_col_map) >= 5:
            data_row_idx = 0
            for ri, row_cells in enumerate(tbl[1:], start=1):
                if len(row_cells) <= max(item_col_map.values(), default=0):
                    continue

                entry: Dict[str, str] = {}
                yn_count = 0
                for item_num, col_idx in item_col_map.items():
                    if col_idx < len(row_cells) and _has_yes_no(row_cells[col_idx]):
                        entry[f"item_{item_num}"] = _normalize_yn(row_cells[col_idx])
                        yn_count += 1

                if yn_count < 3:
                    continue

                if econ_col >= 0 and econ_col < len(row_cells):
                    sm = score_1_5_re.match(row_cells[econ_col].strip())
                    if sm:
                        entry["item_11_economy"] = sm.group(1)

                if prof_col >= 0 and prof_col < len(row_cells):
                    pm = prof_re.search(row_cells[prof_col])
                    if pm:
                        entry["item_12_proficient"] = _normalize_prof(pm.group(1))

                row_id = ""
                if case_col >= 0 and case_col < len(row_cells):
                    cell_val = row_cells[case_col].strip()
                    if cell_val:
                        row_id = cell_val
                if not row_id and data_row_idx < len(candidate_ids):
                    row_id = candidate_ids[data_row_idx]
                if not row_id:
                    row_id = f"docx_row_{ri}"

                surgeon_rows[row_id] = entry
                data_row_idx += 1
            if surgeon_rows:
                continue

        found_items: Dict[int, str] = {}
        tbl_case_id = ""
        item_row_re = re.compile(r"(?:item\s*)?#?\s*(\d{1,2})\b", re.IGNORECASE)

        for row_cells in tbl:
            combined = " ".join(row_cells)
            combined_lower = combined.lower()

            if ("candidate" in combined_lower or "case" in combined_lower) and not tbl_case_id:
                id_m = id_label_pat.search(combined)
                if id_m:
                    tbl_case_id = id_m.group(1).replace(" ", "").replace("-", "_")
                else:
                    short_m = re.search(r"(?:candidate|case)[\s_\-:/#]*(\d{3,6})", combined, re.IGNORECASE)
                    if short_m:
                        tbl_case_id = short_m.group(1)
                continue

            for ci, cell in enumerate(row_cells):
                m = item_row_re.match(cell.strip())
                if not m:
                    continue
                item_num = int(m.group(1))
                if item_num < 1 or item_num > 13:
                    continue

                other_cells = [row_cells[j].strip() for j in range(len(row_cells)) if j != ci and row_cells[j].strip()]

                if item_num <= 10:
                    for oc in other_cells:
                        if _has_yes_no(oc):
                            found_items[item_num] = _normalize_yn(oc)
                            break
                elif item_num == 11:
                    for oc in other_cells:
                        sm = score_1_5_re.match(oc)
                        if sm:
                            found_items[11] = sm.group(1)
                            break
                elif item_num == 12:
                    for oc in other_cells:
                        pm = prof_re.search(oc)
                        if pm:
                            found_items[12] = _normalize_prof(pm.group(1))
                            break

        if len(found_items) >= 5:
            if not tbl_case_id:
                tbl_case_id = candidate_ids[0] if candidate_ids else "docx_case"
            entry = {}
            for item_num, val in found_items.items():
                if item_num <= 10:
                    entry[f"item_{item_num}"] = val
                elif item_num == 11:
                    entry["item_11_economy"] = val
                elif item_num == 12:
                    entry["item_12_proficient"] = val
            surgeon_rows[tbl_case_id] = entry

    if surgeon_rows:
        return surgeon_rows

    paragraphs = [(p.text.strip(), any(r.bold for r in p.runs if r.bold is not None)) for p in doc.paragraphs]
    item_num_re = re.compile(r"^(\d{1,2})\.\s*$")
    rate_re = re.compile(r"[Rr]ate[:\s]*(\d)\s*(?:out\s*of|/)\s*\d")

    found_items_para: Dict[int, str] = {}
    current_item: int = 0
    last_numbered: int = 0
    score_assigned = False

    vop_item_hints: Dict[str, int] = {
        "oblique incision": 1,
        "back wall injury": 2,
        "spatula": 3,
        "double-ended suture": 4,
        "heel of anastomosis": 4,
        "knot on outside": 5,
        "right angle": 6,
        "back wall of anastomosis": 7,
        "toe of graft": 8,
        "front wall of anastomosis": 9,
        "knot on outside of anastomosis": 10,
        "economy of time": 11,
        "demonstrates proficiency": 12,
        "final rating": 12,
        "summative comments": 13,
    }

    for pi, (text, is_bold) in enumerate(paragraphs):
        if not text:
            continue

        m = item_num_re.match(text)
        if m:
            num = int(m.group(1))
            if 1 <= num <= 13:
                current_item = num
                last_numbered = num
                score_assigned = False
                continue

        if current_item == 0:
            text_lower = text.lower()
            for hint, hint_num in vop_item_hints.items():
                if hint in text_lower and hint_num not in found_items_para:
                    if hint_num > last_numbered:
                        current_item = hint_num
                        score_assigned = False
                        break
            continue

        if not score_assigned and current_item <= 10:
            if yes_no_re.match(text):
                found_items_para[current_item] = _normalize_yn(text)
                score_assigned = True
                continue

        if current_item == 11 and not score_assigned:
            rm = rate_re.search(text)
            if rm:
                found_items_para[11] = rm.group(1)
                score_assigned = True
                continue

        if current_item == 12 and not score_assigned:
            if yes_no_re.match(text):
                found_items_para[12] = _normalize_yn(text)
                score_assigned = True
                continue

        if score_assigned and not item_num_re.match(text):
            text_lower = text.lower()
            for hint, hint_num in vop_item_hints.items():
                if hint in text_lower and hint_num not in found_items_para:
                    if hint_num > last_numbered:
                        current_item = hint_num
                        score_assigned = False
                        break

    if len(found_items_para) >= 5:
        case_id = candidate_ids[0] if candidate_ids else "docx_case"
        entry = {}
        for item_num, val in found_items_para.items():
            if item_num <= 10:
                entry[f"item_{item_num}"] = val
            elif item_num == 11:
                entry["item_11_economy"] = val
            elif item_num == 12:
                entry["item_12_proficient"] = val
        surgeon_rows[case_id] = entry

    return surgeon_rows


def build_comparison_analysis(
    batch_results: List[Dict[str, Any]],
    surgeon_rows: Dict[str, Dict[str, str]],
) -> List[Dict[str, Any]]:
    """Build structured comparison between AI and surgeon scores with prose analysis."""
    successful_results = [r for r in batch_results if r.get("success")]
    if not successful_results or not surgeon_rows:
        return []

    all_records = []
    for result in successful_results:
        records = build_result_json(
            result.get("filename", "Unknown"),
            result.get("response", ""),
            result.get("timestamp", ""),
            scored_outputs=result.get("scored_outputs"),
            scoring_warnings=result.get("scoring_warnings"),
        )
        all_records.extend(records)

    def _normalize_key(k: str) -> str:
        return re.sub(r"[\s\-]+", "_", k.strip()).lower()

    norm_surgeon_map: Dict[str, Dict[str, str]] = {}
    for skey, sval in surgeon_rows.items():
        norm_surgeon_map[_normalize_key(skey)] = sval

    comparisons = []
    for record in all_records:
        if record.get("is_legacy"):
            continue
        case_id = record.get("case_id", "")
        candidate_id = record.get("candidate_id", "")
        norm_case = _normalize_key(case_id)
        norm_cand = _normalize_key(candidate_id)

        surgeon_row = (
            surgeon_rows.get(case_id)
            or surgeon_rows.get(candidate_id)
            or norm_surgeon_map.get(norm_case)
            or norm_surgeon_map.get(norm_cand)
        )
        if not surgeon_row:
            for skey, sval in surgeon_rows.items():
                nsk = _normalize_key(skey)
                if candidate_id and (candidate_id in skey or norm_cand in nsk or nsk.endswith(norm_cand)):
                    surgeon_row = sval
                    break
                if case_id and (case_id in skey or norm_case in nsk or nsk == norm_case):
                    surgeon_row = sval
                    break
        if not surgeon_row:
            if len(surgeon_rows) == 1 and len(all_records) == 1:
                surgeon_row = list(surgeon_rows.values())[0]
        if not surgeon_row:
            continue

        items = record.get("items", [])
        item_comparisons = []
        agreements = 0
        disagreements = 0
        nulls = 0
        ai_hawkish = 0
        ai_dovish = 0

        for i in range(1, 11):
            ai_entry = get_item_by_id(items, i)
            ai_score = ai_entry.get("score", "")
            ai_evidence = ai_entry.get("evidence", "")
            surgeon_score = surgeon_row.get(f"item_{i}", "").strip().upper()
            if surgeon_score in ("Y",):
                surgeon_score = "YES"
            elif surgeon_score in ("N",):
                surgeon_score = "NO"

            label = ITEM_LABELS.get(i, f"Item {i}")

            if ai_score == "NULL":
                status = "null"
                nulls += 1
            elif not surgeon_score:
                status = "no_surgeon"
            elif ai_score == surgeon_score:
                status = "agree"
                agreements += 1
            else:
                status = "disagree"
                disagreements += 1
                if ai_score == "NO" and surgeon_score == "YES":
                    ai_hawkish += 1
                elif ai_score == "YES" and surgeon_score == "NO":
                    ai_dovish += 1

            item_comparisons.append({
                "item": i,
                "label": label,
                "ai_score": ai_score,
                "surgeon_score": surgeon_score,
                "status": status,
                "ai_evidence": ai_evidence,
            })

        ai_econ = record.get("economy_score")
        surgeon_econ_str = surgeon_row.get("item_11_economy", "")
        try:
            surgeon_econ = int(surgeon_econ_str)
        except (ValueError, TypeError):
            surgeon_econ = None
        econ_delta = None
        if ai_econ is not None and surgeon_econ is not None:
            econ_delta = ai_econ - surgeon_econ

        ai_prof = record.get("proficiency", "")
        surgeon_prof = surgeon_row.get("item_12_proficient", "").strip().upper()
        if surgeon_prof in ("Y",):
            surgeon_prof = "YES"
        elif surgeon_prof in ("N",):
            surgeon_prof = "NO"

        ai_prof_binary = "YES" if ai_prof == "PROFICIENT" else ("NO" if ai_prof == "NOT_PROFICIENT" else "INSUFFICIENT")
        prof_agree = ai_prof_binary == surgeon_prof if surgeon_prof in ("YES", "NO") else None

        rl = record.get("red_lines_triggered", [])

        scored = agreements + disagreements
        agreement_rate = (agreements / scored * 100) if scored > 0 else 0

        prose_parts = []

        prose_parts.append(f"**Case {case_id}** — AI vs Surgeon Comparison")
        prose_parts.append("")

        if scored > 0:
            prose_parts.append(
                f"Of the {scored} items where both AI and surgeon provided scores, "
                f"they **agreed on {agreements}** ({agreement_rate:.0f}%) and "
                f"**disagreed on {disagreements}**."
            )
        if nulls > 0:
            null_items = [str(ic["item"]) for ic in item_comparisons if ic["status"] == "null"]
            prose_parts.append(
                f"The AI marked {nulls} item(s) as NULL (not observable): Items {', '.join(null_items)}. "
                f"These could not be compared."
            )
        prose_parts.append("")

        if disagreements > 0:
            prose_parts.append("**Disagreements:**")
            for ic in item_comparisons:
                if ic["status"] == "disagree":
                    direction = ""
                    if ic["ai_score"] == "NO" and ic["surgeon_score"] == "YES":
                        direction = " (AI stricter)"
                    elif ic["ai_score"] == "YES" and ic["surgeon_score"] == "NO":
                        direction = " (AI more lenient)"
                    prose_parts.append(
                        f"- **Item {ic['item']}** ({ic['label']}): AI = {ic['ai_score']}, "
                        f"Surgeon = {ic['surgeon_score']}{direction}"
                    )
                    if ic["ai_evidence"]:
                        evidence_short = ic["ai_evidence"][:200]
                        if len(ic["ai_evidence"]) > 200:
                            evidence_short += "..."
                        prose_parts.append(f"  - AI evidence: _{evidence_short}_")
            prose_parts.append("")

        if ai_hawkish > 0 and ai_dovish == 0:
            prose_parts.append(
                f"The AI was **consistently stricter** than the surgeon on {ai_hawkish} item(s). "
                f"This could indicate the AI is being hawkish on these criteria, or the surgeon "
                f"is on the dovish end. Consider reviewing the FOR evidence for these items."
            )
        elif ai_dovish > 0 and ai_hawkish == 0:
            prose_parts.append(
                f"The AI was **consistently more lenient** than the surgeon on {ai_dovish} item(s). "
                f"This could indicate the AI is being dovish, or the surgeon is on the hawkish end."
            )
        elif ai_hawkish > 0 and ai_dovish > 0:
            prose_parts.append(
                f"The AI was stricter on {ai_hawkish} item(s) and more lenient on {ai_dovish}. "
                f"This mixed pattern suggests case-specific differences rather than systematic bias."
            )
        elif disagreements == 0 and scored > 0:
            prose_parts.append(
                "**Full agreement** on all observed checklist items — AI and surgeon are aligned."
            )
        prose_parts.append("")

        if ai_econ is not None and surgeon_econ is not None:
            if econ_delta == 0:
                prose_parts.append(
                    f"**Economy:** AI = {ai_econ}/5, Surgeon = {surgeon_econ} — **Match.**"
                )
            else:
                ed = econ_delta if econ_delta is not None else 0
                direction = "higher" if ed > 0 else "lower"
                prose_parts.append(
                    f"**Economy:** AI = {ai_econ}/5, Surgeon = {surgeon_econ} (delta: {ed:+d}). "
                    f"AI scored {direction} — "
                    f"{'within acceptable range.' if abs(ed) <= 1 else 'notable gap; review economy event detection.'}"
                )
        elif ai_econ is not None:
            prose_parts.append(f"**Economy:** AI = {ai_econ}/5 (no surgeon score to compare).")

        if surgeon_prof in ("YES", "NO"):
            if prof_agree:
                prose_parts.append(
                    f"**Proficiency:** AI = {ai_prof}, Surgeon = {surgeon_prof} — **Match.**"
                )
            else:
                prose_parts.append(
                    f"**Proficiency:** AI = {ai_prof}, Surgeon = {surgeon_prof} — **Mismatch.** "
                )
                if ai_prof == "INSUFFICIENT_EVIDENCE":
                    prose_parts.append(
                        "AI could not determine proficiency due to unobserved core domains. "
                        "This is expected when key items (2, 7, 9, 10) are not visible."
                    )
                elif ai_prof == "NOT_PROFICIENT" and surgeon_prof == "YES":
                    prose_parts.append(
                        f"AI triggered red-line(s): {', '.join(rl) if rl else 'unknown'}. "
                        f"Review whether the AI's NO scores on core items are justified."
                    )
                elif ai_prof == "PROFICIENT" and surgeon_prof == "NO":
                    prose_parts.append(
                        "Surgeon failed the candidate but AI found no red-line triggers. "
                        "The surgeon may be applying criteria not captured in the rubric."
                    )
        prose_parts.append("")

        prose_parts.append("**Tuning Implications:**")
        tuning_notes = []
        if disagreements == 0 and (econ_delta is None or abs(econ_delta) <= 1) and prof_agree is not False:
            tuning_notes.append("No tuning needed — AI and surgeon are well-aligned on this case.")
        else:
            if ai_hawkish > 0:
                hawk_items = [ic for ic in item_comparisons if ic["status"] == "disagree" and ic["ai_score"] == "NO" and ic["surgeon_score"] == "YES"]
                item_nums = [str(ic["item"]) for ic in hawk_items]
                tuning_notes.append(
                    f"Items {', '.join(item_nums)}: AI scored NO where surgeon scored YES. "
                    f"Review FOR extraction — is the AI seeing problems that aren't there, "
                    f"or is the surgeon being lenient?"
                )
            if ai_dovish > 0:
                dove_items = [ic for ic in item_comparisons if ic["status"] == "disagree" and ic["ai_score"] == "YES" and ic["surgeon_score"] == "NO"]
                item_nums = [str(ic["item"]) for ic in dove_items]
                tuning_notes.append(
                    f"Items {', '.join(item_nums)}: AI scored YES where surgeon scored NO. "
                    f"Review whether the AI's FOR captured the relevant observations."
                )
            if econ_delta is not None and abs(econ_delta) > 1:
                tuning_notes.append(
                    f"Economy delta of {econ_delta:+d} exceeds the target range of +/-1. "
                    f"Review wasted-motion event detection thresholds."
                )
            if prof_agree is False and ai_prof == "INSUFFICIENT_EVIDENCE":
                tuning_notes.append(
                    "Proficiency mismatch due to insufficient evidence — "
                    "improving video coverage or camera angle could resolve this."
                )
            if prof_agree is False and ai_prof != "INSUFFICIENT_EVIDENCE":
                tuning_notes.append(
                    "Proficiency mismatch — examine which red-line rules differ "
                    "and whether the AI's item-level scores are justified."
                )

        if not tuning_notes:
            tuning_notes.append("Minor differences within acceptable range. Monitor across more cases.")

        for note in tuning_notes:
            prose_parts.append(f"- {note}")

        comparisons.append({
            "case_id": case_id,
            "item_comparisons": item_comparisons,
            "agreements": agreements,
            "disagreements": disagreements,
            "nulls": nulls,
            "agreement_rate": agreement_rate,
            "ai_hawkish_count": ai_hawkish,
            "ai_dovish_count": ai_dovish,
            "ai_economy": ai_econ,
            "surgeon_economy": surgeon_econ,
            "economy_delta": econ_delta,
            "ai_proficiency": ai_prof,
            "surgeon_proficiency": surgeon_prof,
            "proficiency_agree": prof_agree,
            "red_lines": rl,
            "prose": "\n".join(prose_parts),
        })

    return comparisons


def display_comparison_analysis(comparisons: List[Dict[str, Any]]):
    """Display the AI vs Surgeon comparison inline."""
    if not comparisons:
        return

    st.divider()
    st.subheader("AI vs Surgeon Comparison")

    for comp in comparisons:
        with st.expander(f"Case {comp['case_id']} — {comp['agreements']} agree, {comp['disagreements']} disagree, {comp['nulls']} null", expanded=True):
            agree_pct = comp["agreement_rate"]
            if agree_pct >= 90:
                bar_color = "#4caf50"
            elif agree_pct >= 70:
                bar_color = "#ffc107"
            else:
                bar_color = "#f44336"

            st.markdown(
                f"<div style='background:#222;border-radius:8px;padding:12px;margin-bottom:12px;'>"
                f"<span style='font-size:1.3em;font-weight:bold;color:{bar_color};'>"
                f"Agreement: {agree_pct:.0f}%</span>"
                f"&nbsp;&nbsp;|&nbsp;&nbsp;"
                f"Agree: {comp['agreements']} &nbsp; Disagree: {comp['disagreements']} &nbsp; NULL: {comp['nulls']}"
                f"</div>",
                unsafe_allow_html=True,
            )

            cols = st.columns([1, 2, 1.5, 1.5, 1.5])
            cols[0].markdown("**Item**")
            cols[1].markdown("**Description**")
            cols[2].markdown("**AI**")
            cols[3].markdown("**Surgeon**")
            cols[4].markdown("**Status**")

            for ic in comp["item_comparisons"]:
                cols = st.columns([1, 2, 1.5, 1.5, 1.5])
                cols[0].write(str(ic["item"]))
                cols[1].write(ic["label"])

                ai_color = "#4caf50" if ic["ai_score"] == "YES" else ("#f44336" if ic["ai_score"] == "NO" else "#9e9e9e")
                cols[2].markdown(f"<span style='color:{ai_color};font-weight:bold;'>{ic['ai_score']}</span>", unsafe_allow_html=True)

                s_color = "#4caf50" if ic["surgeon_score"] == "YES" else ("#f44336" if ic["surgeon_score"] == "NO" else "#9e9e9e")
                s_display = ic["surgeon_score"] if ic["surgeon_score"] else "—"
                cols[3].markdown(f"<span style='color:{s_color};font-weight:bold;'>{s_display}</span>", unsafe_allow_html=True)

                if ic["status"] == "agree":
                    status_icon = "<span style='color:#4caf50;'>AGREE</span>"
                elif ic["status"] == "disagree":
                    status_icon = "<span style='color:#f44336;font-weight:bold;'>DIFFER</span>"
                elif ic["status"] == "null":
                    status_icon = "<span style='color:#9e9e9e;'>NULL</span>"
                else:
                    status_icon = "<span style='color:#9e9e9e;'>—</span>"
                cols[4].markdown(status_icon, unsafe_allow_html=True)

            st.divider()

            econ_parts = []
            if comp["ai_economy"] is not None:
                econ_parts.append(f"AI: {comp['ai_economy']}/5")
            if comp["surgeon_economy"] is not None:
                econ_parts.append(f"Surgeon: {comp['surgeon_economy']}")
            if comp["economy_delta"] is not None:
                econ_parts.append(f"Delta: {comp['economy_delta']:+d}")
            if econ_parts:
                econ_color = "#4caf50" if comp["economy_delta"] is not None and abs(comp["economy_delta"]) <= 1 else "#ffc107"
                st.markdown(f"**Economy:** {' | '.join(econ_parts)}")

            prof_parts = [f"AI: {comp['ai_proficiency']}"]
            if comp["surgeon_proficiency"]:
                prof_parts.append(f"Surgeon: {comp['surgeon_proficiency']}")
            if comp["proficiency_agree"] is True:
                prof_parts.append("**Match**")
            elif comp["proficiency_agree"] is False:
                prof_parts.append("**Mismatch**")
            st.markdown(f"**Proficiency:** {' | '.join(prof_parts)}")

            if comp["red_lines"]:
                st.markdown(f"**Red-lines triggered:** {', '.join(comp['red_lines'])}")

            st.divider()
            st.markdown("### Analysis")
            st.markdown(comp["prose"])


def parse_csv_surgeon_scores(csv_text: str) -> Dict[str, Dict[str, str]]:
    """Parse surgeon scores from CSV text into a standardized dict."""
    reader = csv.DictReader(io.StringIO(csv_text))
    surgeon_rows: Dict[str, Dict[str, str]] = {}
    for row in reader:
        key = row.get("case_id", "").strip()
        if not key:
            key = row.get("candidate_id", "").strip()
        if key:
            surgeon_rows[key] = row
    return surgeon_rows


def generate_comparison_csv(batch_results: List[Dict[str, Any]], surgeon_rows: Dict[str, Dict[str, str]]) -> str:
    """Generate surgeon-vs-AI comparison CSV from pre-parsed surgeon rows."""
    successful_results = [r for r in batch_results if r.get("success")]
    if not successful_results or not surgeon_rows:
        return ""

    output = io.StringIO()
    headers = ["case_id"]
    for i in range(1, 11):
        headers += [f"item_{i}_surgeon", f"item_{i}_ai", f"item_{i}_agree"]
    headers += ["economy_surgeon", "economy_ai", "economy_delta",
                 "proficiency_surgeon", "proficiency_ai", "proficiency_agree",
                 "red_lines_ai", "disagreement_notes"]
    writer = csv.writer(output)
    writer.writerow(headers)

    all_records = []
    for result in successful_results:
        records = build_result_json(
            result.get("filename", "Unknown"),
            result.get("response", ""),
            result.get("timestamp", ""),
            scored_outputs=result.get("scored_outputs"),
            scoring_warnings=result.get("scoring_warnings"),
        )
        all_records.extend(records)

    for record in all_records:
        if record.get("is_legacy"):
            continue
        case_id = record.get("case_id", "")
        candidate_id = record.get("candidate_id", "")
        surgeon_row = surgeon_rows.get(case_id) or surgeon_rows.get(candidate_id)
        if not surgeon_row:
            continue

        items = record.get("items", [])
        row = [case_id]

        disagree_notes = []
        for i in range(1, 11):
            ai_entry = get_item_by_id(items, i)
            ai_score = ai_entry.get("score", "")
            surgeon_score = surgeon_row.get(f"item_{i}", "").strip().upper()
            if ai_score == "NULL":
                agree = "NULL"
            else:
                agree = "Y" if ai_score == surgeon_score else "N"
            row += [surgeon_score, ai_score, agree]

            if agree == "N":
                disagree_notes.append(f"Item {i}: AI={ai_score} vs Surgeon={surgeon_score}")
            elif agree == "NULL":
                disagree_notes.append(f"Item {i}: AI=NULL (not observed)")

        ai_econ = record.get("economy_score")
        surgeon_econ = surgeon_row.get("item_11_economy", "")
        try:
            surgeon_econ_int = int(surgeon_econ)
        except (ValueError, TypeError):
            surgeon_econ_int = None
        econ_delta = ""
        if ai_econ is not None and surgeon_econ_int is not None:
            econ_delta = str(ai_econ - surgeon_econ_int)
        row += [surgeon_econ, str(ai_econ) if ai_econ is not None else "", econ_delta]

        ai_prof = record.get("proficiency", "")
        surgeon_prof = surgeon_row.get("item_12_proficient", "").strip().upper()
        prof_agree = "Y" if ai_prof == surgeon_prof else "N"
        row += [surgeon_prof, ai_prof, prof_agree]

        rl = record.get("red_lines_triggered", [])
        row.append("; ".join(rl) if rl else "None")
        row.append("; ".join(disagree_notes) if disagree_notes else "None")
        writer.writerow(row)

    return output.getvalue()


def render_inbox_tab():
    # --- Folder selection ----------------------------------------------------
    folder_input = st.text_input(
        "Video folder path",
        value=st.session_state.get("inbox_folder_path", str(DEFAULT_VIDEOS_INBOX)),
        key="inbox_folder_input",
        help="Point the app at any folder on disk. Subfolders are scanned "
             "recursively. Defaults to the bundled videos_inbox/ folder.",
    )
    folder_input = (folder_input or "").strip()
    if not folder_input:
        folder_input = str(DEFAULT_VIDEOS_INBOX)

    inbox_root = Path(folder_input).expanduser()
    if not inbox_root.exists():
        st.error(f"Folder does not exist: `{inbox_root}`")
        return
    if not inbox_root.is_dir():
        st.error(f"Path is not a directory: `{inbox_root}`")
        return

    # Persist the validated selection for scan/queue/resume.
    st.session_state.inbox_folder_path = str(inbox_root)
    if str(inbox_root) != str(DEFAULT_VIDEOS_INBOX):
        st.caption(f"Using custom folder: `{inbox_root}`")

    with st.expander("Add videos to inbox", expanded=True):
        inbox_upload = st.file_uploader(
            "Upload video to inbox (one at a time)",
            type=["mp4", "mov", "avi", "webm", "mkv", "m4v"],
            accept_multiple_files=False,
            key="inbox_uploader",
            help="Upload a video file (up to 2GB). It will be saved to the selected folder for batch processing.",
        )
        if inbox_upload is not None:
            dest = inbox_root / inbox_upload.name
            if dest.exists():
                st.warning(f"**{inbox_upload.name}** is already in the folder.")
            else:
                with st.spinner(f"Saving {inbox_upload.name} to folder..."):
                    with open(dest, "wb") as f:
                        f.write(inbox_upload.getvalue())
                st.success(f"Saved **{inbox_upload.name}** ({dest.stat().st_size / (1024*1024):.0f} MB) to folder")
                st.rerun()

    inbox_videos = scan_inbox(inbox_root)
    queue = load_queue()
    queue_summary = get_queue_summary(queue)
    has_active_queue = queue.get("session_id", "") and queue_summary["total"] > 0

    if not inbox_videos and not has_active_queue:
        st.info(
            "Upload videos above, or place files directly into the **`videos_inbox/`** folder "
            "(subfolders are scanned recursively). "
            "Videos are processed one at a time from disk with crash recovery."
        )
        return

    if has_active_queue:
        completed = load_completed_results_from_queue(queue)
        pending_names = [
            name for name, info in queue.get("videos", {}).items()
            if info.get("status") in ("pending", "processing")
        ]
        failed_names = [
            name for name, info in queue.get("videos", {}).items()
            if info.get("status") == "failed"
        ]

        st.markdown(
            f"**Queue**: {queue_summary['completed']} completed, "
            f"{queue_summary['pending'] + queue_summary.get('processing', 0)} pending, "
            f"{queue_summary['failed']} failed — out of {queue_summary['total']} total"
        )

        if pending_names:
            with st.expander(f"Pending ({len(pending_names)})"):
                for name in pending_names:
                    st.caption(name)

        if failed_names:
            with st.expander(f"Failed ({len(failed_names)})"):
                for name in failed_names:
                    info = queue["videos"][name]
                    st.caption(f"{name}: {info.get('error', 'unknown error')}")
                if st.button("Retry Failed Videos", key="retry_failed"):
                    for name in failed_names:
                        queue["videos"][name]["status"] = "pending"
                        queue["videos"][name].pop("error", None)
                    save_queue(queue)
                    st.session_state.inbox_processing = True
                    st.rerun()

        new_inbox = [v for v in inbox_videos if inbox_key(v, inbox_root) not in queue.get("videos", {})]
        if new_inbox:
            st.success(f"{len(new_inbox)} new video(s) detected in inbox")
            if st.button("Add New Videos to Queue", key="add_new"):
                for v in new_inbox:
                    queue["videos"][inbox_key(v, inbox_root)] = {
                        "path": str(v),
                        "status": "pending",
                        "added": datetime.now().isoformat(),
                    }
                save_queue(queue)
                st.rerun()

        col_start, col_clear = st.columns(2)
        with col_start:
            if pending_names and not st.session_state.inbox_processing:
                if st.button("Resume Processing", type="primary", use_container_width=True, key="resume_inbox"):
                    st.session_state.inbox_processing = True
                    st.rerun()
        with col_clear:
            if st.button("Clear Queue & Start Fresh", use_container_width=True, key="clear_queue"):
                clear_queue()
                st.session_state.inbox_processing = False
                st.session_state.batch_results = []
                st.rerun()

        if st.session_state.inbox_processing:
            _run_inbox_processing(queue)
            return

        if completed:
            st.session_state.batch_results = completed
            st.divider()
            st.subheader(f"Completed Results ({len(completed)} video{'s' if len(completed) != 1 else ''})")

            is_focused = any(r.get("focused_mode") == "items_1_3" for r in completed if r.get("success"))
            successful = [r for r in completed if r.get("success")]

            if successful:
                file_stem = build_batch_filename_stem(completed)
                col1, col2, col3 = st.columns(3)
                with col1:
                    batch_json = generate_batch_json(completed)
                    if batch_json:
                        st.download_button(
                            "Download All JSON",
                            data=batch_json,
                            file_name=f"vop_results_{file_stem}.json",
                            mime="application/json",
                            key="download_json_inbox",
                        )
                with col2:
                    if is_focused:
                        csv_data = generate_focused_csv(completed)
                    else:
                        csv_data = generate_scores_csv(completed)
                    if csv_data:
                        label = "Download Items 1&3 CSV" if is_focused else "Download CSV"
                        st.download_button(
                            label,
                            data=csv_data,
                            file_name=f"vop_scores_{file_stem}.csv",
                            mime="text/csv",
                            key="download_csv_inbox",
                        )
                with col3:
                    if not is_focused:
                        pdf_zip = generate_all_pdfs_zip(completed)
                        if pdf_zip:
                            st.download_button(
                                "Download all PDFs (ZIP)",
                                data=pdf_zip,
                                file_name=f"vop_reports_{file_stem}.zip",
                                mime="application/zip",
                                key="download_pdfs_inbox",
                            )

            for result in completed:
                if is_focused:
                    display_focused_result(result)
                else:
                    display_assessment_result(result)

            if st.session_state.surgeon_scores and not is_focused:
                comparisons = build_comparison_analysis(completed, st.session_state.surgeon_scores)
                if comparisons:
                    display_comparison_analysis(comparisons)

    else:
        st.success(f"**{len(inbox_videos)} video(s)** found in inbox (subfolders scanned recursively)")
        for v in inbox_videos:
            size_mb = v.stat().st_size / (1024 * 1024)
            st.caption(f"{inbox_key(v, inbox_root)} ({size_mb:.0f} MB)")

        if st.button(
            f"Start Processing ({len(inbox_videos)} video{'s' if len(inbox_videos) > 1 else ''})",
            type="primary",
            use_container_width=True,
            key="start_inbox",
        ):
            queue = init_or_resume_queue(inbox_videos, st.session_state.focused_mode)
            st.session_state.inbox_processing = True
            st.rerun()


def _run_inbox_processing(queue: Dict[str, Any]):
    session_id = queue["session_id"]
    focused = queue.get("focused_mode", st.session_state.focused_mode)

    pending = [
        (name, info) for name, info in queue.get("videos", {}).items()
        if info.get("status") in ("pending", "processing")
    ]

    if not pending:
        st.session_state.inbox_processing = False
        failed_count = sum(1 for info in queue["videos"].values() if info["status"] == "failed")
        if failed_count:
            st.warning(f"Processing finished — {failed_count} video(s) failed. Use 'Retry Failed' to reprocess them.")
        else:
            st.success("All videos processed successfully!")
        st.rerun()
        return

    total = len(queue.get("videos", {}))
    completed_so_far = sum(1 for info in queue["videos"].values() if info["status"] == "completed")

    name, info = pending[0]
    video_path = info["path"]

    if not Path(video_path).exists():
        queue["videos"][name]["status"] = "failed"
        queue["videos"][name]["error"] = "File not found in inbox"
        save_queue(queue)
        st.rerun()
        return

    progress_text = f"Processing {completed_so_far + 1} of {total}: {name}"
    st.progress(completed_so_far / total, text=progress_text)

    queue["videos"][name]["status"] = "processing"
    queue["videos"][name]["started"] = datetime.now().isoformat()
    save_queue(queue)

    with st.spinner(f"Extracting observations from {name}..."):
        client = GeminiVisionClient()
        # Pass the basename to the client (candidate-ID / video-id parsing works
        # on the filename), but keep the relative key `name` for queue bookkeeping.
        result = process_video_from_path(video_path, Path(name).name, client, focused_mode=focused)

        safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
        if result.get("success"):
            save_result_to_disk(result, session_id)
            queue["videos"][name]["status"] = "completed"
            queue["videos"][name]["completed"] = datetime.now().isoformat()
            queue["videos"][name]["result_file"] = f"{safe_name}.json"
            save_queue(queue)
        else:
            queue["videos"][name]["status"] = "failed"
            queue["videos"][name]["error"] = result.get("error", "Unknown error")
            queue["videos"][name]["completed"] = datetime.now().isoformat()
            save_queue(queue)

        st.rerun()


def main():
    init_session_state()

    st.title("Vascular Anastomosis VOP Assessment")
    st.caption(
        "Retrospective Proficiency Assessment — Official 2023 VOP Rubric (Tri-State Scoring + Coverage)"
    )

    with st.sidebar:
        st.header("Configuration")

        input_mode = st.radio(
            "Video Input",
            ["Upload", "Inbox Folder"],
            horizontal=True,
            help="**Upload**: select files through browser (best for 1-2 videos). "
                 "**Inbox Folder**: place videos in the `videos_inbox/` folder (subfolders scanned recursively) "
                 "for robust batch processing with crash recovery.",
        )

        uploaded_files = []
        if input_mode == "Upload":
            uploaded_files = st.file_uploader(
                "Upload Anastomosis Video(s)",
                type=["mp4", "mov", "avi", "webm", "mkv", "m4v"],
                accept_multiple_files=True,
                help=f"Upload video(s) of simulated vascular anastomosis (max {MAX_FILE_SIZE_MB}MB per file). Videos over ~20 minutes are automatically split for processing (3 fps sampling).",
            )

            if uploaded_files:
                st.info(f"**{len(uploaded_files)} video(s) selected**")
                for f in uploaded_files:
                    size_mb = len(f.getvalue()) / (1024 * 1024)
                    status = "ok" if size_mb <= MAX_FILE_SIZE_MB else "too large"
                    st.caption(f"{f.name} ({size_mb:.1f}MB) — {status}")

        st.session_state.focused_mode = st.toggle(
            "Items 1 & 3 Only",
            value=st.session_state.focused_mode,
            help="Score only Item 1 (arteriotomy incision) and Item 3 (graft spatulation). Faster processing, lower quota usage.",
        )
        if st.session_state.focused_mode:
            st.caption("Focused mode: only scoring arteriotomy incision and graft spatulation")

        surgeon_assessments = st.file_uploader(
            "Upload Surgeon Assessment(s) (optional)",
            type=["csv", "docx"],
            accept_multiple_files=True,
            key="surgeon_assessment_files",
            help="Upload one or more surgeon-scored CSV or DOCX files to automatically compare with AI results after processing.",
        )

        if surgeon_assessments:
            merged_scores: Dict[str, Dict[str, str]] = {}
            filenames = []
            for sa_file in surgeon_assessments:
                sa_bytes = sa_file.getvalue()
                sa_name = sa_file.name.lower()
                filenames.append(sa_file.name)
                try:
                    if sa_name.endswith(".docx"):
                        parsed = parse_docx_surgeon_scores(sa_bytes, filename=sa_file.name)
                    else:
                        parsed = parse_csv_surgeon_scores(sa_bytes.decode("utf-8"))
                    if parsed:
                        merged_scores.update(parsed)
                    else:
                        st.warning(f"Could not extract scores from: {sa_file.name}")
                except Exception as e:
                    st.warning(f"Error parsing {sa_file.name}: {e}")
            if merged_scores:
                st.session_state.surgeon_scores = merged_scores
                st.session_state.surgeon_filename = ", ".join(filenames)
                st.success(
                    f"Surgeon scores loaded from {len(surgeon_assessments)} file(s): "
                    f"{len(merged_scores)} candidate(s) — {', '.join(merged_scores.keys())}"
                )
            else:
                st.session_state.surgeon_scores = {}
        else:
            if not st.session_state.surgeon_scores:
                st.caption("Optional: upload surgeon assessment(s) for automatic comparison")

        st.divider()

        with st.expander("About the Rubric (2023 VOP)"):
            st.markdown(
                """
**VOP 2023 v1 — Evidence-Based Scoring:**

**Stage 1** — Factual Observation Record (FOR):
Gemini extracts purely descriptive observations from the video.

**Stage 2** — LLM Scoring with Validators:
Gemini scores from the FOR JSON. App-side validators enforce schema, economy guardrails, and proficiency override.

**Tri-State Scoring (Items 1-10):** YES / NO / NULL
- OBSERVED: directly seen in video
- NOT_OBSERVED: not visible → score = NULL (no assumptions)

**Item 6 Split:** 6a (right-angle method, non-gating) + 6b (safe transfer outcome). Overall = 6b score.

**Coverage Metrics:** observed_count, observed_percent, core_observed (Items 2, 7, 9, 10)

**3-State Proficiency:**
- PROFICIENT: no red-lines, all core domains observed
- NOT_PROFICIENT: any core domain NO
- INSUFFICIENT_EVIDENCE: any core domain NULL

**Red-Line Rules:**
- RL-A: Item 7 = NO → fails
- RL-B: Item 9 = NO → fails
- RL-C: Item 2 = NO → fails
- RL-D: Item 10 = NO → fails
- RL-E: Item 4 = NO (observed) → fails

**Economy Index (out of 5):** Weighted wasted-motion events, event-based (no per-minute normalization)

**Projection Layer:** hawk_mode / construct_mode / dove_mode for DOCX-compatible binary output
            """
            )

        st.divider()

        with st.expander("Surgeon Comparison"):
            st.markdown("Upload surgeon-scored **CSV** or **DOCX** to compare AI vs surgeon scores. NULL items show as 'NULL' in agreement column.")
            surgeon_file = st.file_uploader(
                "Upload Surgeon Scores",
                type=["csv", "docx"],
                key="surgeon_file",
                help="CSV (columns: case_id, item_1...item_10, item_11_economy, item_12_proficient) or DOCX scoring form",
            )
            if surgeon_file and st.session_state.batch_results:
                file_bytes = surgeon_file.getvalue()
                file_name = surgeon_file.name.lower()
                surgeon_rows: Dict[str, Dict[str, str]] = {}
                parse_error = ""

                try:
                    if file_name.endswith(".docx"):
                        surgeon_rows = parse_docx_surgeon_scores(file_bytes, filename=surgeon_file.name)
                        if surgeon_rows:
                            st.info(f"Parsed {len(surgeon_rows)} candidate(s) from DOCX: {', '.join(surgeon_rows.keys())}")
                        else:
                            parse_error = "Could not extract surgeon scores from DOCX. Ensure the file contains a table with item scores (YES/NO) and candidate IDs."
                    else:
                        csv_text = file_bytes.decode("utf-8")
                        surgeon_rows = parse_csv_surgeon_scores(csv_text)
                        if not surgeon_rows:
                            parse_error = "No valid rows found in CSV. Ensure it has a case_id or candidate_id column."
                except Exception as e:
                    parse_error = f"Error parsing file: {e}"

                if parse_error:
                    st.warning(parse_error)
                elif surgeon_rows:
                    comparison = generate_comparison_csv(st.session_state.batch_results, surgeon_rows)
                    if comparison:
                        st.download_button(
                            "Download Comparison CSV",
                            data=comparison,
                            file_name=f"surgeon_vs_ai_{build_batch_filename_stem(st.session_state.batch_results)}.csv",
                            mime="text/csv",
                            key="comparison_csv",
                        )
                        st.success("Comparison generated!")
                    else:
                        st.warning("No matching cases found between surgeon scores and AI results.")

        st.divider()
        st.info("Powered by **Gemini 3.1 Pro** — 2023 VOP Rubric (Tri-State + Coverage)")
        st.caption("Videos up to 1GB. Videos over ~20 min auto-split (3 fps sampling). Evidence-based scoring with app-side validators.")

        saved_sessions = list_saved_sessions()
        if saved_sessions:
            st.divider()
            st.subheader("Saved Sessions")
            for session in saved_sessions[:5]:
                session_label = (
                    f"{session['session_id']} ({session['file_count']} videos)"
                )
                if st.button(
                    f"Load: {session_label}", key=f"load_{session['session_id']}"
                ):
                    st.session_state.batch_results = load_session_results(
                        session["session_id"]
                    )
                    st.session_state.session_id = session["session_id"]
                    st.rerun()

    if input_mode == "Inbox Folder":
        render_inbox_tab()
        return

    if not uploaded_files:
        st.info("Upload one or more simulated vascular anastomosis videos to begin assessment")

        with st.expander("About this tool"):
            st.markdown(
                """
This tool assesses videos of **simulated vascular anastomosis** performed by
first-year medical residents using the **Moderated v1** architecture.

**How it works:**
1. Upload one or more videos (or use **Inbox Folder** mode for large batches)
2. Click "Start Assessment"
3. **Stage 1**: Gemini extracts a Factual Observation Record (FOR) — purely descriptive
4. **Stage 2**: Gemini scores from FOR with app-side validators:
   - **Tri-state scoring**: YES / NO / NULL for 10 procedural items
   - **Coverage metrics**: observed_count, observed_percent, core_observed
   - **Economy Index (out of 5)**: weighted events, event-based (no per-minute normalization)
   - **3-state proficiency**: PROFICIENT / NOT_PROFICIENT / INSUFFICIENT_EVIDENCE
   - **Projection layer**: hawk / construct / dove modes for DOCX output

**Key policies (VOP 2023 v1):**
- **No assumptions**: Unobserved items score NULL (not defaulted to YES)
- **Core domains**: Items 2, 7, 9, 10 — any NULL triggers INSUFFICIENT_EVIDENCE
- **Item 6 split**: 6a (method) + 6b (outcome); overall = 6b; non-gating
- **Economy Index**: Weighted events (instrument_search=2, failed_pass=2, regrasp=1)
- **Projection layer**: Converts tri-state to binary for DOCX compatibility

**Video Requirements:**
- Maximum 1GB per file (videos over ~20 minutes automatically split into parts; sampled at 3 fps)
- Supported formats: MP4, MOV, AVI, WebM, MKV, M4V

**Batch Processing:**
- For large batches (many videos or large files), use **Inbox Folder** mode
- Videos are processed one at a time from disk — no memory issues
- Results are saved after each video — crash-safe with automatic resume
            """
            )

        if st.session_state.batch_results:
            st.divider()
            st.subheader("Previous Results")

            successful_results = [
                r for r in st.session_state.batch_results if r.get("success")
            ]
            loaded_is_focused = any(r.get("focused_mode") == "items_1_3" for r in st.session_state.batch_results if r.get("success"))

            if successful_results:
                col1, col2, col3 = st.columns(3)

                file_stem = build_batch_filename_stem(st.session_state.batch_results)

                with col1:
                    batch_json = generate_batch_json(st.session_state.batch_results)
                    if batch_json:
                        st.download_button(
                            "Download All JSON",
                            data=batch_json,
                            file_name=f"vop_results_{file_stem}.json",
                            mime="application/json",
                            key="download_json_loaded",
                        )

                with col2:
                    if loaded_is_focused:
                        csv_data = generate_focused_csv(st.session_state.batch_results)
                    else:
                        csv_data = generate_scores_csv(st.session_state.batch_results)
                    if csv_data:
                        label = "Download Items 1&3 CSV" if loaded_is_focused else "Download CSV"
                        st.download_button(
                            label,
                            data=csv_data,
                            file_name=f"vop_scores_{file_stem}.csv",
                            mime="text/csv",
                            key="download_csv_loaded",
                        )

                with col3:
                    if not loaded_is_focused:
                        pdf_zip = generate_all_pdfs_zip(st.session_state.batch_results)
                        if pdf_zip:
                            st.download_button(
                                "Download all PDFs (ZIP)",
                                data=pdf_zip,
                                file_name=f"vop_reports_{file_stem}.zip",
                                mime="application/zip",
                                key="download_pdfs_loaded",
                            )

            for result in st.session_state.batch_results:
                if loaded_is_focused:
                    display_focused_result(result)
                else:
                    display_assessment_result(result)

            if st.session_state.surgeon_scores and not loaded_is_focused:
                comparisons = build_comparison_analysis(
                    st.session_state.batch_results,
                    st.session_state.surgeon_scores,
                )
                if comparisons:
                    display_comparison_analysis(comparisons)
                else:
                    st.divider()
                    st.warning(
                        f"Surgeon assessment loaded ({st.session_state.surgeon_filename}) "
                        f"but no matching cases found. Surgeon cases: "
                        f"{', '.join(st.session_state.surgeon_scores.keys())}"
                    )

        return

    session_id = get_session_id()

    if st.session_state.processing:
        idx = st.session_state.current_index
        total = st.session_state.total_files
        pending = st.session_state.pending_files

        if idx < total:
            current_file = pending[idx]
            progress_text = f"Processing video {idx + 1} of {total}: {current_file.name}"
            progress_bar = st.progress((idx) / total, text=progress_text)

            with st.spinner(f"Extracting observations from {current_file.name}..."):
                try:
                    client = GeminiVisionClient()
                    result = process_single_video(current_file, client, focused_mode=st.session_state.focused_mode)

                    st.session_state.batch_results.append(result)

                    if result.get("success"):
                        save_result_to_disk(result, session_id)

                    st.session_state.current_index = idx + 1
                    progress_bar.progress(
                        (idx + 1) / total,
                        text=f"Completed {idx + 1} of {total}",
                    )
                    st.rerun()
                except Exception as e:
                    st.error(f"Error processing {current_file.name}: {str(e)}")
                    error_result = {
                        "filename": current_file.name,
                        "success": False,
                        "error": str(e),
                        "timestamp": datetime.now().isoformat(),
                    }
                    st.session_state.batch_results.append(error_result)
                    st.session_state.current_index = idx + 1
                    st.rerun()
        else:
            st.session_state.processing = False
            st.rerun()

    else:
        col1, col2 = st.columns([3, 1])
        with col1:
            if st.button(
                f"Start Assessment ({len(uploaded_files)} video{'s' if len(uploaded_files) > 1 else ''})",
                type="primary",
                use_container_width=True,
            ):
                st.session_state.batch_results = []
                st.session_state.pending_files = uploaded_files
                st.session_state.current_index = 0
                st.session_state.total_files = len(uploaded_files)
                st.session_state.processing = True
                st.session_state.session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
                st.rerun()

    if st.session_state.batch_results and not st.session_state.processing:
        st.divider()
        successful = [r for r in st.session_state.batch_results if r.get("success")]
        failed = [r for r in st.session_state.batch_results if not r.get("success")]
        total = len(st.session_state.batch_results)

        st.subheader(
            f"Assessment Results: {len(successful)} of {total} successful"
        )

        is_focused = any(r.get("focused_mode") == "items_1_3" for r in st.session_state.batch_results if r.get("success"))

        if successful:
            file_stem = build_batch_filename_stem(st.session_state.batch_results)
            col1, col2 = st.columns(2)
            with col1:
                batch_json = generate_batch_json(st.session_state.batch_results)
                if batch_json:
                    st.download_button(
                        "Download All JSON",
                        data=batch_json,
                        file_name=f"vop_results_{file_stem}.json",
                        mime="application/json",
                        key="download_all_batch",
                    )
            with col2:
                if is_focused:
                    csv_data = generate_focused_csv(st.session_state.batch_results)
                else:
                    csv_data = generate_scores_csv(st.session_state.batch_results)
                if csv_data:
                    label = "Download Items 1&3 CSV" if is_focused else "Download CSV"
                    st.download_button(
                        label,
                        data=csv_data,
                        file_name=f"vop_scores_{file_stem}.csv",
                        mime="text/csv",
                        key="download_csv_batch",
                    )

        for result in st.session_state.batch_results:
            if is_focused:
                display_focused_result(result)
            else:
                display_assessment_result(result)

        if st.session_state.surgeon_scores and not is_focused:
            comparisons = build_comparison_analysis(
                st.session_state.batch_results,
                st.session_state.surgeon_scores,
            )
            if comparisons:
                display_comparison_analysis(comparisons)
            else:
                st.divider()
                st.warning(
                    f"Surgeon assessment loaded ({st.session_state.surgeon_filename}) "
                    f"but no matching cases found. Surgeon cases: "
                    f"{', '.join(st.session_state.surgeon_scores.keys())}"
                )


if __name__ == "__main__":
    main()
